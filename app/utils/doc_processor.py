"""
Document processing utilities — upgraded to use AWS Textract + Comprehend.

When USE_TEXTRACT=true (production):
  - Textract extracts text and tables from PDF/image files stored in S3.
  - Comprehend detects entities (ORG, PERSON, QUANTITY) for vendor/price extraction.

When USE_TEXTRACT=false (local dev):
  - Falls back to pdfminer / PyPDF2 / python-docx (original behaviour).
"""

import os
import re
import time
import logging
from typing import Dict, Any

logger = logging.getLogger(__name__)


# ─────────────────────────── Entry point ─────────────────────────────────────

def extract_document_data(file_path: str, use_textract: bool = None) -> Dict[str, Any]:
    """
    Detect file type and extract structured data.
    file_path is an S3 key (when USE_S3=true) or a local absolute path.
    """
    from flask import current_app
    if use_textract is None:
        use_textract = current_app.config.get('USE_TEXTRACT', False)

    if not file_path:
        return {}

    ext = file_path.rsplit('.', 1)[-1].lower() if '.' in file_path else ''

    try:
        if use_textract and ext in ('pdf', 'png', 'jpg', 'jpeg'):
            return _extract_via_textract(file_path)
        elif ext == 'pdf':
            return _extract_from_pdf(file_path)
        elif ext in ('docx', 'doc'):
            return _extract_from_docx(file_path)
        else:
            logger.info(f'No extractor for extension: {ext}')
            return {}
    except Exception as e:
        logger.error(f'Document extraction error ({file_path}): {e}')
        return {'extraction_error': str(e)}


# ─────────────────────────── AWS Textract ────────────────────────────────────

def _extract_via_textract(s3_key: str) -> Dict[str, Any]:
    """
    Use Amazon Textract to extract text + tables from a document in S3.
    Uses start_document_analysis (async) for multi-page PDFs.
    """
    import boto3
    from flask import current_app

    region = current_app.config.get('AWS_S3_REGION', 'ap-south-1')
    bucket = current_app.config['AWS_S3_BUCKET']

    textract = boto3.client('textract', region_name=region)

    # Start async job (supports multi-page PDFs)
    response = textract.start_document_analysis(
        DocumentLocation={'S3Object': {'Bucket': bucket, 'Name': s3_key}},
        FeatureTypes=['TABLES', 'FORMS'],
    )
    job_id = response['JobId']
    logger.info(f'Textract job started: {job_id} for {s3_key}')

    # Poll until complete (max 120 seconds)
    text_blocks = []
    table_cells = []
    for _ in range(24):
        time.sleep(5)
        result = textract.get_document_analysis(JobId=job_id)
        status = result['JobStatus']
        if status == 'SUCCEEDED':
            text_blocks, table_cells = _parse_textract_blocks(result['Blocks'])
            # Handle pagination
            while 'NextToken' in result:
                result = textract.get_document_analysis(
                    JobId=job_id, NextToken=result['NextToken'])
                tb, tc = _parse_textract_blocks(result['Blocks'])
                text_blocks.extend(tb)
                table_cells.extend(tc)
            break
        elif status == 'FAILED':
            logger.error(f'Textract job failed: {job_id}')
            return {'extraction_error': 'Textract job failed'}

    full_text = ' '.join(text_blocks)

    # Use Comprehend for entity extraction
    entities = _extract_entities_comprehend(full_text, region)

    result = _parse_extracted_text(full_text, source='textract')
    result['textract_job_id'] = job_id
    result['comprehend_entities'] = entities

    # Override vendor name with Comprehend ORG entity if found
    org_entities = [e['Text'] for e in entities if e['Type'] == 'ORGANIZATION']
    if org_entities and not result.get('vendor_name'):
        result['vendor_name'] = org_entities[0]

    # Extract prices from table cells
    table_prices = []
    for cell in table_cells:
        amt = _parse_currency(cell)
        if amt and amt > 1000:  # Filter out small numbers
            table_prices.append(amt)
    if table_prices:
        result.setdefault('prices_found', [])
        result['prices_found'] = sorted(
            set(result['prices_found'] + table_prices), reverse=True)

    return result


def _parse_textract_blocks(blocks):
    """Extract LINE text and CELL text from Textract block list."""
    text_lines = []
    cell_texts = []
    for block in blocks:
        if block['BlockType'] == 'LINE':
            text_lines.append(block.get('Text', ''))
        elif block['BlockType'] == 'CELL':
            text = block.get('Text', '')
            if text:
                cell_texts.append(text)
    return text_lines, cell_texts


def _extract_entities_comprehend(text: str, region: str) -> list:
    """
    Use Amazon Comprehend to detect entities.
    Returns list of {Type, Text, Score} dicts.
    Comprehend has a 5000 UTF-8 byte limit per call.
    """
    try:
        import boto3
        comprehend = boto3.client('comprehend', region_name=region)
        # Truncate to Comprehend limit
        chunk = text[:4900]
        response = comprehend.detect_entities(Text=chunk, LanguageCode='en')
        return response.get('Entities', [])
    except Exception as e:
        logger.warning(f'Comprehend entity detection failed: {e}')
        return []


# ─────────────────────────── Local PDF / DOCX ─────────────────────────────────

def _extract_from_pdf(file_path: str) -> Dict[str, Any]:
    text = ''
    try:
        from pdfminer.high_level import extract_text as pdfminer_extract
        text = pdfminer_extract(file_path) or ''
    except Exception as e:
        logger.debug(f'pdfminer failed, trying PyPDF2: {e}')

    if not text.strip():
        try:
            import PyPDF2
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                pages = [page.extract_text() or '' for page in reader.pages]
                text = '\n'.join(pages)
        except Exception as e:
            logger.error(f'PyPDF2 extraction failed: {e}')

    return _parse_extracted_text(text, source='pdf')


def _extract_from_docx(file_path: str) -> Dict[str, Any]:
    try:
        from docx import Document
        doc = Document(file_path)
        para_text = '\n'.join(p.text for p in doc.paragraphs)
        table_data = []
        price_from_tables = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(cells)
                for cell in cells:
                    amt = _parse_currency(cell)
                    if amt:
                        price_from_tables.append(amt)
            table_data.append(rows)
        result = _parse_extracted_text(para_text, source='docx')
        result['raw_tables'] = table_data
        if price_from_tables and not result.get('prices_found'):
            result['prices_found'] = price_from_tables
        return result
    except Exception as e:
        logger.error(f'DOCX extraction error: {e}')
        return {'extraction_error': str(e)}


# ─────────────────────────── Parsing helpers ──────────────────────────────────

def _parse_extracted_text(text: str, source: str) -> Dict[str, Any]:
    result: Dict[str, Any] = {
        'source': source,
        'char_count': len(text),
        'prices_found': [],
        'vendor_name': None,
        'gst_number': None,
        'pan_number': None,
        'emails': [],
        'phones': [],
    }
    if not text.strip():
        return result

    price_patterns = [
        r'(?:₹|Rs\.?|INR)\s*([\d,]+(?:\.\d{1,2})?)',
        r'(?:total|amount|price|cost|value|bid)\s*:?\s*(?:₹|Rs\.?|INR)?\s*([\d,]+(?:\.\d{1,2})?)',
    ]
    prices = []
    for pat in price_patterns:
        for match in re.finditer(pat, text, re.IGNORECASE):
            raw = match.group(1).replace(',', '')
            try:
                prices.append(float(raw))
            except ValueError:
                pass
    result['prices_found'] = sorted(set(prices), reverse=True)

    gst_match = re.search(
        r'\b(\d{2}[A-Z]{5}\d{4}[A-Z]{1}[1-9A-Z]{1}Z[0-9A-Z]{1})\b', text)
    if gst_match:
        result['gst_number'] = gst_match.group(1)

    pan_match = re.search(r'\b([A-Z]{5}[0-9]{4}[A-Z]{1})\b', text)
    if pan_match:
        result['pan_number'] = pan_match.group(1)

    emails = re.findall(
        r'\b[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}\b', text)
    result['emails'] = list(set(emails))[:5]

    phones = re.findall(r'\b(?:\+91[\-\s]?)?[6-9]\d{9}\b', text)
    result['phones'] = list(set(phones))[:5]

    company_match = re.search(
        r'(?:company|firm|vendor|bidder|name)\s*:?\s*([A-Z][A-Za-z\s&\.,]{3,60})',
        text, re.IGNORECASE)
    if company_match:
        result['vendor_name'] = company_match.group(1).strip()

    return result


def _parse_currency(text: str):
    text = text.replace(',', '').strip()
    match = re.search(r'(?:₹|Rs\.?|INR)?\s*([\d]+(?:\.\d{1,2})?)', text)
    if match:
        try:
            val = float(match.group(1))
            if val > 0:
                return val
        except ValueError:
            pass
    return None
